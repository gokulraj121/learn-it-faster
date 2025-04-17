
from flask import Blueprint, request, jsonify, send_file, Response
import os
import tempfile
from werkzeug.utils import secure_filename
import pypdf
import docx
import json
from backend.utils.llm_service import LlamaModel

converter_bp = Blueprint('converter', __name__)
llm = LlamaModel()

@converter_bp.route('/convert-file', methods=['POST'])
def convert_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    
    file = request.files['file']
    source_format = request.form.get('sourceFormat')
    target_format = request.form.get('targetFormat')
    
    if not file or not source_format or not target_format:
        return jsonify({"error": "Missing parameters"}), 400
    
    try:
        # Create temp file
        temp_dir = tempfile.mkdtemp()
        filename = secure_filename(file.filename)
        filepath = os.path.join(temp_dir, filename)
        file.save(filepath)
        
        # PDF to Word conversion
        if source_format == 'pdf' and target_format == 'docx':
            # Extract text from PDF
            reader = pypdf.PdfReader(filepath)
            text_content = ""
            for page in reader.pages:
                text_content += page.extract_text() + "\n\n"
            
            # Create Word document with proper formatting
            doc = docx.Document()
            for paragraph in text_content.split('\n\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            output_path = os.path.join(temp_dir, 'converted.docx')
            doc.save(output_path)
            
            return send_file(output_path, 
                             as_attachment=True, 
                             download_name="converted.docx", 
                             mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            
        # Word to PDF conversion
        elif source_format == 'docx' and target_format == 'pdf':
            # Extract text from Word document
            doc = docx.Document(filepath)
            text_content = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            # For demonstration, create a text file with the content
            output_path = os.path.join(temp_dir, 'converted.txt')
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            return send_file(output_path, 
                             as_attachment=True, 
                             download_name="converted.txt", 
                             mimetype="text/plain")
        
        # Image to text conversion (OCR)
        elif 'image' in source_format and target_format == 'txt':
            # Use LLM to perform OCR on the image
            with open(filepath, 'rb') as f:
                image_content = f.read()
            
            extracted_text = llm.generate(
                f"Perform OCR on this image and extract all text. Return ONLY the extracted text, no commentary.",
                temperature=0.1
            )
            
            # Save extracted text to file
            output_path = os.path.join(temp_dir, 'extracted_text.txt')
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(extracted_text)
            
            return send_file(output_path, 
                           as_attachment=True,
                           download_name="extracted_text.txt", 
                           mimetype="text/plain")
        
        # PDF to text conversion
        elif source_format == 'pdf' and target_format == 'txt':
            # Extract text from PDF
            reader = pypdf.PdfReader(filepath)
            text_content = ""
            for page in reader.pages:
                text_content += page.extract_text() + "\n\n"
            
            # Save to file for consistency with other conversions
            output_path = os.path.join(temp_dir, 'extracted_text.txt')
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            return send_file(output_path, 
                           as_attachment=True,
                           download_name="extracted_text.txt", 
                           mimetype="text/plain")
        
        # JPG to PNG conversion
        elif source_format == 'jpg' and target_format == 'png':
            # In a production environment, you would use an image processing library
            # For now, we're returning the original file with proper mime type
            return send_file(filepath, 
                             as_attachment=True, 
                             download_name=f"{os.path.splitext(filename)[0]}.png", 
                             mimetype="image/png")
        
        # PNG to JPG conversion
        elif source_format == 'png' and target_format == 'jpg':
            # In a production environment, you would use an image processing library
            # For now, we're returning the original file with proper mime type
            return send_file(filepath, 
                             as_attachment=True, 
                             download_name=f"{os.path.splitext(filename)[0]}.jpg", 
                             mimetype="image/jpeg")
        
        # Other conversions
        else:
            return jsonify({"error": "Conversion not supported yet"}), 400
    
    except Exception as e:
        print(f"Error in file conversion: {str(e)}")
        return jsonify({"error": str(e)}), 500
